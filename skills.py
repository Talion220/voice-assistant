import os
import webbrowser
import sys
import subprocess
import datetime
import threading
import queue
import json
import sounddevice as sd
import vosk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import voice

try:
    import requests
except:
    pass

recording_active = False
recording_thread = None
recording_queue = queue.Queue()


def browser():
    webbrowser.open('https://ya.ru', new=2)


def game():
    try:
        subprocess.Popen('C:/Windows/system32/mspaint.exe')
    except:
        voice.speaker('Путь к файлу не найден, проверьте, правильный ли он')


def time():
    now = datetime.datetime.now()
    hours = now.hour
    minutes = now.minute
    voice.speaker(f"Текущее время: {hours} часов {minutes} минут")


def offpc():
    #отключает ПК под управлением Windows

    #os.system('shutdown \s')
    print('пк был бы выключен, но команде # в коде мешает;)))')


def weather():
    '''Для работы нужно зарегистрироваться на сайте https://openweathermap.org '''
    try:
        params = {'q': 'Krasnoyarsk', 'units': 'metric', 'lang': 'ru', 'appid': 'f5b76d1ea4b8a6b22819e0fd530612ec'}
        response = requests.get(f'https://api.openweathermap.org/data/2.5/weather', params=params)
        if not response:
            raise
        w = response.json()
        voice.speaker(f"На улице {w['weather'][0]['description']} {round(w['main']['temp'])} градусов")
    except:
        voice.speaker('Произошла ошибка при попытке запроса к ресурсу API, проверь код')


def offBot():
    voice.speaker("Пока")
    os._exit(0)


def start_recording():
    global recording_active, recording_thread, recording_queue

    if recording_active:
        voice.speaker("Запись уже идет")
        return

    recording_active = True
    recording_queue = queue.Queue()

    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = os.path.join(desktop_path, f"Диктовка_{timestamp}.docx")

    voice.speaker("Начинаю запись. Скажите стоп запись для завершения")

    recording_thread = threading.Thread(target=_record_to_word, args=(filename,), daemon=True)
    recording_thread.start()


def stop_recording():
    global recording_active

    if not recording_active:
        return

    recording_active = False
    voice.speaker("Запись остановлена")


def _record_to_word(filename):
    global recording_active, recording_queue

    model = vosk.Model('model_small')
    device = sd.default.device
    samplerate = int(sd.query_devices(device[0], 'input')['default_samplerate'])

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    title = doc.add_heading('Диктовка', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(f"Дата: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}").italic = True

    doc.add_paragraph()

    full_text = []
    current_paragraph = ""
    stop_phrases = ['стоп запись', 'остановить запись', 'закончить запись', 'хватит записывать']

    def audio_callback(indata, frames, time, status):
        if recording_active:
            recording_queue.put(bytes(indata))

    rec = vosk.KaldiRecognizer(model, samplerate)

    try:
        with sd.RawInputStream(
                samplerate=samplerate,
                blocksize=16000,
                device=device[0],
                dtype='int16',
                channels=1,
                callback=audio_callback
        ):
            while recording_active:
                try:
                    data = recording_queue.get(timeout=0.5)
                except queue.Empty:
                    continue

                if rec.AcceptWaveform(data):
                    result = json.loads(rec.Result())
                    text = result.get('text', '').strip()

                    if text:
                        text_lower = text.lower()

                        stop_recording_now = False
                        for phrase in stop_phrases:
                            if phrase in text_lower:
                                stop_recording_now = True
                                break

                        if stop_recording_now:
                            if current_paragraph:
                                full_text.append(current_paragraph.strip())
                                current_paragraph = ""
                            break

                        if text.endswith(('.', '!', '?')):
                            current_paragraph += " " + text
                            full_text.append(current_paragraph.strip())
                            current_paragraph = ""
                        else:
                            if current_paragraph:
                                current_paragraph += " " + text
                            else:
                                current_paragraph = text

        if current_paragraph.strip():
            full_text.append(current_paragraph.strip())

        for paragraph_text in full_text:
            doc.add_paragraph(paragraph_text)

        doc.save(filename)
        voice.speaker(f"Файл сохранен на рабочем столе")

    except Exception as e:
        print(f"Ошибка записи: {e}")
        voice.speaker("Произошла ошибка при записи")

    finally:
        recording_active = False


def passive():
    '''Функция заглушка при простом диалоге с ботом'''
    pass

